"""
Paper Writer Agent

Generates the Research Starter Pack paper draft — section by section,
with every AI-generated section clearly watermarked.

Design principles:
    1. One Claude call per section (keeps context focused, reduces hallucination)
    2. Every section grounded in real data from ExecutionResults and DataHealthReport
    3. Every AI-generated section marked [DRAFT — RESEARCHER REVIEW NEEDED]
    4. Placeholders [RESEARCHER TO FILL] where only the researcher can contribute
    5. Structured skeleton, not a finished paper

The paper draft is a STARTING POINT. AutoResearch runs baselines.
The researcher's novel contribution comes after.

Sections generated:
    - abstract          (grounded in real results)
    - introduction      (problem motivation — researcher fills domain expertise)
    - related_work      (Semantic Scholar citations — researcher validates)
    - methodology       (describes what was actually run)
    - experiments       (real numbers from ExecutionResults)
    - results           (comparison table + plain English interpretation)
    - discussion        (honest analysis of what worked and why)
    - conclusion        (summary + future work suggestions)
    - limitations       (explicit about what AutoResearch can't do)
"""

import json
import logging
import re
from pathlib import Path
from typing import Optional

from autoresearch.agents.api_utils import LLMClient

from autoresearch.schemas import (
    ProblemSpec,
    DataHealthReport,
    MethodsCatalog,
    ExecutionResult,
    EvaluationReport,
    ExecutionStatus,
)

logger = logging.getLogger(__name__)

# Watermark added to every AI-generated section
DRAFT_WATERMARK = "[DRAFT — RESEARCHER REVIEW NEEDED]"
RESEARCHER_FILL = "[RESEARCHER TO FILL]"


# ── Section prompts ───────────────────────────────────────────────────────────

SECTION_SYSTEM_PROMPT = """\
You are writing a section of a research paper draft for AutoResearch.

Critical rules:
1. Write in academic style but keep it readable
2. ONLY state things that are supported by the data provided
3. Never invent results, citations, or claims not in the input
4. Use [RESEARCHER TO FILL] for anything requiring domain expertise or
   personal knowledge the researcher must provide
5. This is a DRAFT — be honest about what AutoResearch did and didn't do
6. Cite methods by name and results by their actual numbers
7. Write concisely — researchers will expand sections they care about

Output: plain text only. No markdown headers (caller adds those).
"""


def build_abstract_prompt(
    spec: ProblemSpec,
    report: DataHealthReport,
    evaluation: EvaluationReport,
    results: list[ExecutionResult],
) -> str:
    winner_result = next(
        (r for r in results if r.method_id == evaluation.winner_method_id), None
    )
    metric_str = (
        f"{winner_result.primary_metric_name} = {winner_result.primary_metric_value:.4f}"
        if winner_result and winner_result.primary_metric_value
        else "metric not available"
    )

    return f"""Write a 150-200 word abstract for a research paper.

Problem: {spec.plain_english_summary}
Domain: {spec.domain}
Dataset: {report.row_count:,} rows, {report.column_count} features
Methods tried: {', '.join(r.method_name for r in results if r.status == ExecutionStatus.SUCCESS)}
Best method: {evaluation.winner_method_id} ({metric_str})

The abstract should cover: motivation, data, methods, best result, and key finding.
End with: "Code and reproducible experiments are available at [RESEARCHER TO FILL]."

Write the abstract text only. No title, no labels.
"""


def build_introduction_prompt(spec: ProblemSpec, report: DataHealthReport) -> str:
    return f"""Write a 300-400 word introduction for a research paper.

Problem domain: {spec.domain}
Task type: {spec.task_type.value}
Dataset size: {report.row_count:,} rows
Primary metric: {spec.primary_metric.value}

The introduction should:
1. Open with why this problem matters (use [RESEARCHER TO FILL] for specific clinical/economic/social impact)
2. Describe the challenge briefly
3. State what this paper contributes
4. Outline the paper structure

Include at least 2 [RESEARCHER TO FILL] placeholders where the researcher must add:
- Domain-specific motivation (e.g. patient outcomes, economic cost)
- Specific claims about the research gap this addresses

Write the section text only.
"""


def build_related_work_prompt(
    spec: ProblemSpec,
    papers: list[dict],
) -> str:
    papers_str = "\n".join([
        f"- {p.get('title', 'Unknown')} ({p.get('year', '?')}): {p.get('abstract_snippet', '')}"
        for p in papers[:8]
    ]) if papers else "No papers retrieved — researcher must add citations manually."

    return f"""Write a 300-400 word related work section.

Domain: {spec.domain}
Task type: {spec.task_type.value}

Potentially relevant papers (retrieved from Semantic Scholar — researcher must validate):
{papers_str}

Rules:
- Only cite papers from the list above — do NOT invent citations
- Use (Author et al., Year) format
- Note gaps in the literature that this work addresses
- Include: "[RESEARCHER TO FILL: Add 3-5 papers most relevant to your specific dataset/setting]"
- Include: "[RESEARCHER TO FILL: Describe how your approach differs from prior work]"

Write the section text only.
"""


def build_methodology_prompt(
    spec: ProblemSpec,
    report: DataHealthReport,
    catalog: MethodsCatalog,
) -> str:
    methods_list = "\n".join([
        f"- {m.name}: {m.why_chosen}"
        for m in catalog.methods
    ])

    return f"""Write a 400-500 word methodology section describing the experimental setup.

Dataset:
- {report.row_count:,} rows, {report.column_count} columns
- Task: {spec.task_type.value}
- Target: {spec.target_column or 'none (unsupervised)'}
- Primary metric: {spec.primary_metric.value}
- Data health score: {report.health_score:.0f}/100
- Key data notes: {'; '.join(report.recommendations[:2]) if report.recommendations else 'none'}

Methods evaluated:
{methods_list}

Experimental setup:
- All experiments run on Kaggle Kernels (free tier: T4 GPU / 2-core CPU)
- Hyperparameter optimization: Optuna (30 trials per method)
- Validation: Cross-validation

Include:
- Brief description of each method and why it was chosen
- Feature engineering steps applied
- How hyperparameters were selected
- "[RESEARCHER TO FILL: Describe any domain-specific preprocessing you added]"

Write the section text only.
"""


def build_experiments_prompt(
    spec: ProblemSpec,
    results: list[ExecutionResult],
) -> str:
    successful = [r for r in results if r.status == ExecutionStatus.SUCCESS]
    results_table = "\n".join([
        f"- {r.method_name}: {r.primary_metric_name}={r.primary_metric_value:.4f}"
        f"{f', train={r.train_metric:.4f}' if r.train_metric else ''}"
        f"{f', val={r.val_metric:.4f}' if r.val_metric else ''}"
        f"{f', {r.runtime_minutes:.1f} min' if r.runtime_minutes else ''}"
        for r in successful
    ])
    failed = [r for r in results if r.status != ExecutionStatus.SUCCESS]

    return f"""Write a 200-300 word experiments section.

Primary metric: {spec.primary_metric.value}

Results (real numbers from actual runs):
{results_table}

Failed experiments: {', '.join(r.method_name for r in failed) if failed else 'none'}

Describe:
1. The experimental conditions (hardware, framework)
2. Present results factually — these are real numbers
3. Note any failed experiments and why (if known)

Do NOT interpret results here — that's the Results section.
Write the section text only.
"""


def build_results_prompt(
    spec: ProblemSpec,
    results: list[ExecutionResult],
    evaluation: EvaluationReport,
) -> str:
    successful = [r for r in results if r.status == ExecutionStatus.SUCCESS]
    scores_str = "\n".join([
        f"- {s.method_name}: total_score={s.total_score:.3f}, "
        f"perf={s.score_performance:.3f}, "
        f"interpretability={s.score_interpretability:.3f}"
        for s in evaluation.method_scores
    ])

    return f"""Write a 350-450 word results section.

Winner: {evaluation.winner_method_id}
Winner explanation: {evaluation.winner_explanation}

Multi-criteria scores:
{scores_str}

Risk flags detected:
{chr(10).join(f"- [{f.severity}] {f.explanation}" for f in evaluation.risk_flags) or 'none'}

Include:
1. A sentence describing the comparison table (actual numbers)
2. Analysis of why the winner outperformed others
3. Honest discussion of any risk flags (overfitting, suspicious scores)
4. Note limitations: "Results reflect performance on the provided dataset and
   may not generalize to other settings without further validation."

Write the section text only.
"""


def build_discussion_prompt(
    spec: ProblemSpec,
    evaluation: EvaluationReport,
) -> str:
    return f"""Write a 300-400 word discussion section.

Domain: {spec.domain}
Winner: {evaluation.winner_method_id}
Winner explanation: {evaluation.winner_explanation}
Research directions suggested: {'; '.join(evaluation.research_directions[:3]) if evaluation.research_directions else 'none'}
Constraints: {'; '.join(f"{c.name} ({'required' if c.is_hard else 'preferred'})" for c in spec.constraints) if spec.constraints else 'none'}

Include:
1. Why the winning method worked for this problem
2. What the results imply for the domain (use [RESEARCHER TO FILL] for domain-specific implications)
3. Unexpected findings (if any risk flags were raised)
4. "[RESEARCHER TO FILL: Discuss what these results mean for practitioners in your field]"

Write the section text only.
"""


def build_conclusion_prompt(
    spec: ProblemSpec,
    evaluation: EvaluationReport,
) -> str:
    return f"""Write a 200-250 word conclusion.

Problem: {spec.plain_english_summary}
Best result: {evaluation.winner_method_id} — {evaluation.winner_explanation[:200]}
Suggested next steps: {'; '.join(evaluation.next_steps[:3]) if evaluation.next_steps else 'none'}

Include:
1. One-paragraph summary of what was done and found
2. 2-3 concrete future work directions
3. "[RESEARCHER TO FILL: Add the specific contribution to your field that you will pursue next]"

Write the section text only.
"""


def build_limitations_prompt(
    spec: ProblemSpec,
    report: DataHealthReport,
    results: list[ExecutionResult],
) -> str:
    failed_count = sum(1 for r in results if r.status != ExecutionStatus.SUCCESS)

    return f"""Write a 200-250 word limitations section. Be honest.

Known limitations of this study:
- Dataset: {report.row_count:,} rows (may be small for generalizable conclusions)
- Data health score: {report.health_score:.0f}/100 (issues: {'; '.join(r[2][:80] for r in report.flags[:2]) if report.flags else 'none major'})
- Failed experiments: {failed_count}
- Hyperparameter search: 30 Optuna trials (not exhaustive)
- Compute: Kaggle free tier (GPU quota limited)
- Paper draft: AI-generated skeleton, not peer-reviewed analysis

Be direct about:
1. What the dataset size means for statistical power
2. What hyperparameter search coverage means
3. That this is a baseline study — novel contributions require further work
4. That all AI-generated sections require expert review

Write the section text only.
"""


# ── Paper Writer ──────────────────────────────────────────────────────────────

class PaperWriterAgent:
    """
    The Paper Writer agent.

    Generates each section separately with focused prompts, grounded
    exclusively in real data from the run. Watermarks every AI-generated
    section. Exports to .docx (and optionally .tex).

    Usage:
        agent = PaperWriterAgent(api_key="...")
        sections = agent.write(spec, report, catalog, results, evaluation)
        agent.export_docx(sections, output_dir)
    """

    SECTIONS = [
        "abstract",
        "introduction",
        "related_work",
        "methodology",
        "experiments",
        "results",
        "discussion",
        "conclusion",
        "limitations",
    ]

    def __init__(self, llm: LLMClient, verbose: bool = True):
        self.llm     = llm
        self.verbose = verbose

    def write(
        self,
        spec: ProblemSpec,
        report: DataHealthReport,
        catalog: MethodsCatalog,
        results: list[ExecutionResult],
        evaluation: EvaluationReport,
        papers: Optional[list[dict]] = None,
    ) -> dict[str, str]:
        """
        Write all paper sections. Returns dict of section_name → text.
        Each section is watermarked as a draft.
        """
        if self.verbose:
            print("\n" + "━" * 60)
            print("📄 WRITING PAPER DRAFT")
            print("━" * 60)
            print("\n  Writing each section separately...")
            print("  Every section will be marked as [DRAFT — RESEARCHER REVIEW NEEDED]\n")

        sections: dict[str, str] = {}

        prompts = {
            "abstract":     build_abstract_prompt(spec, report, evaluation, results),
            "introduction": build_introduction_prompt(spec, report),
            "related_work": build_related_work_prompt(spec, papers or []),
            "methodology":  build_methodology_prompt(spec, report, catalog),
            "experiments":  build_experiments_prompt(spec, results),
            "results":      build_results_prompt(spec, results, evaluation),
            "discussion":   build_discussion_prompt(spec, evaluation),
            "conclusion":   build_conclusion_prompt(spec, evaluation),
            "limitations":  build_limitations_prompt(spec, report, results),
        }

        for section_name in self.SECTIONS:
            if self.verbose:
                print(f"  Writing {section_name}...", end=" ", flush=True)

            text = self._write_section(section_name, prompts[section_name])
            sections[section_name] = f"{DRAFT_WATERMARK}\n\n{text}"

            if self.verbose:
                word_count = len(text.split())
                print(f"✓ ({word_count} words)")

        if self.verbose:
            total_words = sum(len(s.split()) for s in sections.values())
            print(f"\n  ✅ Draft complete: {total_words:,} words across {len(sections)} sections")
            print(f"  Remember: every section needs your review before submission.")
            print("━" * 60 + "\n")

        return sections

    def _write_section(self, section_name: str, prompt: str) -> str:
        """Write a single section via Claude."""
        return self.llm.create(
            system=SECTION_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            verbose=self.verbose,
        ).strip()

    def export_docx(
        self,
        sections: dict[str, str],
        spec: ProblemSpec,
        evaluation: EvaluationReport,
        results: list[ExecutionResult],
        output_dir: Path,
    ) -> Path:
        """
        Export paper draft to .docx format.
        Returns path to the created file.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for Word export. "
                "Run: pip install python-docx"
            )

        doc = Document()

        # Title
        title = doc.add_heading(
            f"Research Draft: {spec.domain.title()} ({spec.task_type.value.title()})",
            level=0
        )

        # AutoResearch notice
        notice = doc.add_paragraph()
        notice.add_run(
            "⚠️  AutoResearch Draft — All AI-generated sections require expert review. "
            "Sections marked [RESEARCHER TO FILL] require your input before submission."
        ).bold = True
        notice.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Red

        doc.add_paragraph()

        # Results summary table
        doc.add_heading("Experiment Results Summary", level=2)
        successful = [r for r in results if r.status == ExecutionStatus.SUCCESS]
        if successful:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Method"
            hdr[1].text = "Primary Metric"
            hdr[2].text = "Train Score"
            hdr[3].text = "Val Score"
            for r in successful:
                row = table.add_row().cells
                row[0].text = r.method_name
                row[1].text = f"{r.primary_metric_value:.4f}" if r.primary_metric_value else "—"
                row[2].text = f"{r.train_metric:.4f}" if r.train_metric else "—"
                row[3].text = f"{r.val_metric:.4f}" if r.val_metric else "—"

        doc.add_paragraph()

        # Each section
        section_titles = {
            "abstract":     "Abstract",
            "introduction": "1. Introduction",
            "related_work": "2. Related Work",
            "methodology":  "3. Methodology",
            "experiments":  "4. Experimental Setup",
            "results":      "5. Results",
            "discussion":   "6. Discussion",
            "conclusion":   "7. Conclusion",
            "limitations":  "8. Limitations",
        }

        for section_name, title in section_titles.items():
            text = sections.get(section_name, "")
            if not text:
                continue

            doc.add_heading(title, level=1)

            # Split watermark from content
            if DRAFT_WATERMARK in text:
                watermark_para = doc.add_paragraph()
                run = watermark_para.add_run(DRAFT_WATERMARK)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)  # Orange
                text = text.replace(DRAFT_WATERMARK, "").strip()

            # Add content paragraph by paragraph
            for para_text in text.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                para = doc.add_paragraph()
                # Highlight [RESEARCHER TO FILL] in red
                if RESEARCHER_FILL in para_text:
                    parts = para_text.split(RESEARCHER_FILL)
                    for j, part in enumerate(parts):
                        if part:
                            para.add_run(part)
                        if j < len(parts) - 1:
                            run = para.add_run(RESEARCHER_FILL)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                else:
                    para.add_run(para_text)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / "paper_draft.docx"
        doc.save(str(output_path))

        if self.verbose:
            print(f"  📄 Word document saved: {output_path}")

        return output_path

    def export_markdown(
        self,
        sections: dict[str, str],
        spec: ProblemSpec,
        output_dir: Path,
    ) -> Path:
        """Export paper draft to Markdown (always generated — no dependencies)."""
        section_titles = {
            "abstract":     "Abstract",
            "introduction": "Introduction",
            "related_work": "Related Work",
            "methodology":  "Methodology",
            "experiments":  "Experimental Setup",
            "results":      "Results",
            "discussion":   "Discussion",
            "conclusion":   "Conclusion",
            "limitations":  "Limitations",
        }

        lines = [
            f"# Research Draft: {spec.domain.title()}",
            f"**Task:** {spec.task_type.value} | "
            f"**Metric:** {spec.primary_metric.value}",
            "",
            "> ⚠️ AutoResearch Draft — All sections marked "
            f"`{DRAFT_WATERMARK}` require expert review.",
            "> Sections marked `[RESEARCHER TO FILL]` require your input.",
            "",
        ]

        for section_name, title in section_titles.items():
            text = sections.get(section_name, "")
            if text:
                lines.append(f"## {title}")
                lines.append("")
                lines.append(text)
                lines.append("")

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / "paper_draft.md"
        output_path.write_text("\n".join(lines))

        if self.verbose:
            print(f"  📝 Markdown draft saved: {output_path}")

        return output_path
